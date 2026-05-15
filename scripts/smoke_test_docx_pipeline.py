#!/usr/bin/env python3
"""Smoke test for Module A DOCX/raw extraction → audit → clean JSON pipeline."""
from __future__ import annotations

import json
import re
import subprocess
import sys
import tempfile
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
AUDIT = ROOT / "scripts" / "audit_reading_source.py"
NORMALIZE = ROOT / "scripts" / "normalize_reading_json.py"
VALIDATE = ROOT / "scripts" / "validate_reading.py"
GEN = ROOT / "scripts" / "generate_reading.py"
CLEAN_RAW = ROOT / "fixtures" / "docx_extraction" / "mock_raw_extraction_clean.json"
CONFLICT_RAW = ROOT / "fixtures" / "source_audit" / "mock_raw_extraction_conflict.json"
QUESTION_PATTERNS_RAW = ROOT / "fixtures" / "docx_extraction" / "mock_raw_extraction_question_patterns.json"

BANNED_INTERNAL_PHRASES = ["extracted from teacher answer key", "marked option", "source trace", "raw extraction", "normalizer fallback"]
BANNED_VISIBLE_MARKS = ["\u2713", "\u2714", "\u221a", "[correct]", "(correct)", "[FALSE]", "[TRUE]", "[NOT GIVEN]", "[YES]", "[NO]"]


def run(cmd):
    return subprocess.run(cmd, cwd=ROOT, text=True, capture_output=True)


def main() -> int:
    print("DOCX PIPELINE SMOKE TEST")
    errors = []
    with tempfile.TemporaryDirectory(prefix="reading-docx-pipeline-") as td:
        td = Path(td)
        # Clean fixture should audit and normalize to schema-compatible JSON.
        audit_json = td / "clean_audit.json"
        audit_md = td / "clean_audit.md"
        clean_json = td / "clean_candidate.json"
        html = td / "clean_candidate.html"
        r = run([sys.executable, "-B", str(AUDIT), "--raw", str(CLEAN_RAW), "--json-output", str(audit_json), "--md-output", str(audit_md), "--clean-json", str(clean_json), "--html", str(html)])
        print(r.stdout.strip())
        if r.returncode:
            print(r.stderr); errors.append("clean audit failed")
        else:
            data = json.loads(audit_json.read_text())
            if data.get("overall_verdict") != "PASS":
                errors.append(f"clean fixture expected PASS, got {data.get('overall_verdict')}")
        r = run([sys.executable, "-B", str(NORMALIZE), "--raw", str(CLEAN_RAW), "--audit", str(audit_json), "--output", str(clean_json)])
        print(r.stdout.strip())
        if r.returncode:
            print(r.stderr); errors.append("clean normalize failed")
        r = run([sys.executable, "-B", str(GEN), "--data", str(clean_json), "--output", str(html)])
        print(r.stdout.strip())
        if r.returncode:
            print(r.stderr); errors.append("clean generated HTML failed")
        r = run([sys.executable, "-B", str(VALIDATE), str(html)])
        print(r.stdout.strip())
        if r.returncode:
            print(r.stderr); errors.append("clean generated HTML validation failed")

        # Synthetic U03-style DOCX: Reading Passage N — Title + Level: Band meta, no END separators.
        # This specifically guards the generalized passage detection fallback.
        docx_path = td / "u03_style_no_end.docx"
        doc = Document()
        for n, title, band, genre in [(1, "First New Format", "6.0–6.5", "Analytical Exposition"), (2, "Second New Format", "6.5–7.0", "Social Study")]:
            doc.add_paragraph(f"Reading Passage {n} — {title}")
            doc.add_paragraph(f"Level: Band {band} · {genre}")
            doc.add_paragraph(f"[A] Paragraph A for passage {n}.")
            doc.add_paragraph(f"[B] Paragraph B for passage {n}.")
            doc.add_paragraph("Questions")
            doc.add_paragraph(f"{n}. A direct test statement. [TRUE]")
        doc.save(docx_path)
        raw_u03 = td / "u03_style_raw.json"
        r = run([sys.executable, "-B", str(ROOT / "scripts" / "extract_docx_to_reading_json.py"), "--docx", str(docx_path), "--output", str(raw_u03)])
        print(r.stdout.strip())
        if r.returncode:
            print(r.stderr); errors.append("u03-style extraction command failed")
        else:
            raw_data = json.loads(raw_u03.read_text())
            if len(raw_data.get("passages", [])) != 2:
                errors.append(f"u03-style fixture expected 2 passages, got {len(raw_data.get('passages', []))}")
            metas = [(p.get("band"), p.get("genre")) for p in raw_data.get("passages", [])]
            if not metas or metas[0][0] != "6.0–6.5" or metas[0][1] != "Analytical Exposition":
                errors.append(f"u03-style fixture meta parse failed: {metas}")

        # Conflict fixture must detect answer_key_conflict + summary mismatch.
        audit_json2 = td / "conflict_audit.json"
        audit_md2 = td / "conflict_audit.md"
        r = run([sys.executable, "-B", str(AUDIT), "--raw", str(CONFLICT_RAW), "--json-output", str(audit_json2), "--md-output", str(audit_md2)])
        print(r.stdout.strip())
        if r.returncode:
            print(r.stderr); errors.append("conflict audit command failed")
        else:
            data = json.loads(audit_json2.read_text())
            issue_types = {i["issue_type"] for i in data.get("issues", [])}
            if "answer_key_conflict" not in issue_types:
                errors.append("conflict fixture did not detect answer_key_conflict")
            if "summary_answer_mismatch" not in issue_types:
                errors.append("conflict fixture did not detect summary_answer_mismatch")
            if data.get("overall_verdict") == "PASS":
                errors.append("conflict fixture unexpectedly PASS")

        # Regression: question patterns fixture covers MCQ leakage, TFNG inline, summary multi-blank,
        # matching select, heading matching, and reveal fallback.
        raw_qp = QUESTION_PATTERNS_RAW
        audit_json3 = td / "qp_audit.json"
        audit_md3 = td / "qp_audit.md"
        clean_json3 = td / "qp_clean.json"
        html3 = td / "qp.html"
        r = run([sys.executable, "-B", str(AUDIT), "--raw", str(raw_qp), "--json-output", str(audit_json3), "--md-output", str(audit_md3), "--clean-json", str(clean_json3), "--html", str(html3)])
        print(r.stdout.strip())
        if r.returncode:
            print(r.stderr); errors.append("question patterns audit failed")
        r = run([sys.executable, "-B", str(NORMALIZE), "--raw", str(raw_qp), "--audit", str(audit_json3), "--output", str(clean_json3)])
        print(r.stdout.strip())
        if r.returncode:
            print(r.stderr); errors.append("question patterns normalize failed")
        else:
            clean = json.loads(clean_json3.read_text())
            passages = clean.get("passages", [])
            # Check question counts
            atomic_total = 0
            for p in passages:
                qs = []
                for b in p.get("questions", {}).get("blocks", []):
                    if b.get("type") == "heading_info":
                        qs.append(b.get("q"))
                    elif b.get("type") == "summary":
                        for pair in b.get("pairs", []):
                            qs.append(pair[0])
                            qs += [pair[i] for i in range(4, len(pair) - 1, 2) if isinstance(pair[i], int)]
                    else:
                        qs.append(b.get("q"))
                atomic_total += len(set(qs))
            if atomic_total < 14:
                errors.append(f"question patterns fixture: expected >=14 atomic questions, got {atomic_total}")
            # Check MCQ options don't contain banned visible marks
            for p in passages:
                for b in p.get("questions", {}).get("blocks", []):
                    if b.get("type") == "mcq":
                        for opt in b.get("options", []):
                            for mark in BANNED_VISIBLE_MARKS:
                                if mark in opt:
                                    errors.append(f"MCQ option still contains visible mark {mark!r}: {opt[:60]}")
                    if b.get("type") in ("tfng", "ynng"):
                        stmt = b.get("statement", "")
                        for mark in BANNED_VISIBLE_MARKS:
                            if mark in stmt:
                                errors.append(f"TFNG/YNNG statement still contains mark {mark!r}: {stmt[:60]}")
                    # Check reveal does not contain banned internal phrases
                    reveal = b.get("reveal", "")
                    for phrase in BANNED_INTERNAL_PHRASES:
                        if phrase in reveal.lower():
                            errors.append(f"Reveal contains banned phrase {phrase!r}")
            # Check heading passage has proper structure
            for p in passages:
                if p.get("questions", {}).get("type") == "heading":
                    if not p.get("questions", {}).get("headings"):
                        errors.append("heading passage missing headings list")
                    if not p.get("questions", {}).get("answers"):
                        errors.append("heading passage missing answers map")
        r = run([sys.executable, "-B", str(GEN), "--data", str(clean_json3), "--output", str(html3)])
        print(r.stdout.strip())
        if r.returncode:
            print(r.stderr); errors.append("QP html generation failed")
        else:
            raw_html = html3.read_text()
            # Visible leakage check on generated HTML
            for mark in BANNED_VISIBLE_MARKS:
                # Only check student-visible area (body but not data-ans, not answer-reveal)
                visible = raw_html
                # Remove hidden reveal blocks from visible scan
                visible = re.sub(r'<div class="answer-reveal[^>]*>.*?</div>', '', visible, flags=re.S)
                if mark in visible:
                    errors.append(f"Generated HTML visible area contains leakage marker {mark!r}")

    if errors:
        print("DOCX PIPELINE SMOKE TEST FAILED ❌")
        for e in errors:
            print(" -", e)
        return 1
    print("DOCX PIPELINE SMOKE TEST PASSED ✅")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
